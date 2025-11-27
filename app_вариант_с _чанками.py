# -*- coding: utf-8 -*-
"""
Локальный Flask-сервер для Q&A по загруженным пользователем документам:
- Загрузка файлов (PDF, DOCX, XLSX, TXT) в рамках пользовательской сессии
- Извлечение текста и формирование кэша
- Вопрос → поиск/анализ строго по кэшу
- Поиск ускорен за счёт векторного индекса (эмбеддинги + семантический поиск)
- Вывод ответа + экспорт результата в XLSX

Зависимости: flask, requests, python-docx, PyPDF2, openpyxl, numpy
"""
import os
import io
import json
import uuid
import shutil
import zipfile
from pathlib import Path
from typing import Dict, Any, List, Tuple

from flask import Flask, request, jsonify, send_from_directory, send_file
import requests
from docx import Document
import PyPDF2
from openpyxl import load_workbook as xl_load
from openpyxl.workbook import Workbook

import numpy as np  # ✅ для векторного поиска

# ===================== НАСТРОЙКИ =====================
app = Flask(__name__, static_folder='.')
app.config["MAX_CONTENT_LENGTH"] = 512 * 1024 * 1024  # 512MB

XAI_API_KEY = os.environ.get("XAI_API_KEY")
API_URL = "https://api.x.ai/v1/chat/completions"
MODEL = "grok-4-1-fast-reasoning-latest"
OUTPUT_TOKENS = 10000

# Эмбеддинги (проверь модель/endpoint по своей фактической доке xAI)
EMBEDDING_API_URL = "https://api.x.ai/v1/embeddings"
EMBEDDING_MODEL = "embedding-beta"

# Длина текстового чанка для индекса (~2k символов)
CHUNK_SIZE = 2000

# Сессии (in-memory).
# "chunks": List[str]        — текстовые чанки
# "embeddings": List[List]   — соответствующие векторы
SESSIONS: Dict[str, Dict[str, Any]] = {}


# ===================== УТИЛИТЫ: СЕССИИ =====================
def _ensure_session(session_id: str) -> Dict[str, Any]:
    if not session_id:
        session_id = str(uuid.uuid4())
    if session_id not in SESSIONS:
        base = Path("/tmp") / f"session_{session_id}"
        base.mkdir(parents=True, exist_ok=True)
        SESSIONS[session_id] = {
            "dir": str(base),
            "cache": "",
            "files": [],
            "last_answer": "",
            "last_question": "",
            "chunks": [],        # ✅ для RAG
            "embeddings": []     # ✅ для RAG
        }
    return SESSIONS[session_id]


def _is_allowed(fname: str) -> bool:
    low = fname.lower()
    return (
        low.endswith(".pdf")
        or low.endswith(".docx")
        or low.endswith(".xlsx")
        or low.endswith(".txt")
        or low.endswith(".zip")
    )


# ===================== УТИЛИТЫ: ИЗВЛЕЧЕНИЕ ТЕКСТА =====================
def _extract_text_from_pdf(path: str) -> str:
    try:
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            return "\n".join([page.extract_text() or "" for page in reader.pages])
    except Exception as e:
        return f"[PDF error: {e}]"


def _extract_text_from_docx(path: str) -> str:
    try:
        doc = Document(path)
        parts = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        # простая выборка из таблиц (если есть)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if row_text:
                    parts.append(" | ".join(row_text))
        return "\n".join(parts)
    except Exception as e:
        return f"[DOCX error: {e}]"


def _extract_text_from_xlsx(path: str) -> str:
    """
    Плоское извлечение текста из всех листов XLSX, по строкам.
    Формат: "=== SHEET: Name ===", далее строки "A1 | A2 | ...".
    """
    try:
        wb = xl_load(filename=path, read_only=True, data_only=True)
        out_lines = []
        for ws in wb.worksheets:
            out_lines.append(f"=== SHEET: {ws.title} ===")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip() != ""]
                if cells:
                    out_lines.append(" | ".join(cells))
        return "\n".join(out_lines)
    except Exception as e:
        return f"[XLSX error: {e}]"


def _extract_text_from_txt(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()
    except Exception as e:
        return f"[TXT error: {e}]"


def _ingest_file_to_cache(session: Dict[str, Any], file_path: str, shown_name: str) -> None:
    text = ""
    low = shown_name.lower()
    if low.endswith(".pdf"):
        text = _extract_text_from_pdf(file_path)
    elif low.endswith(".docx"):
        text = _extract_text_from_docx(file_path)
    elif low.endswith(".xlsx"):
        text = _extract_text_from_xlsx(file_path)
    elif low.endswith(".txt"):
        text = _extract_text_from_txt(file_path)

    if text:
        header = f"\n\n### FILE: {shown_name}\n"
        session["cache"] += header + text

    # Важно: после добавления новых файлов старый векторный индекс уже невалиден
    session["chunks"] = []
    session["embeddings"] = []


# ===================== УТИЛИТЫ: ЭМБЕДДИНГИ И RAG =====================
def _embed_texts(texts: List[str]) -> List[List[float]]:
    """
    Возвращает список векторов для списка строк.
    Формат ответа предполагается openai-совместимый:
    {
      "data": [
        {"embedding": [...], ...},
        ...
      ]
    }
    """
    if not texts:
        return []

    if not XAI_API_KEY:
        raise RuntimeError("XAI_API_KEY не задан")

    headers = {
        "Authorization": f"Bearer {XAI_API_KEY}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": EMBEDDING_MODEL,
        "input": texts,
    }
    resp = requests.post(EMBEDDING_API_URL, headers=headers, json=payload, timeout=60)
    resp.raise_for_status()
    data = resp.json()
    return [item["embedding"] for item in data.get("data", [])]


def _split_text_into_chunks(text: str, max_chars: int = CHUNK_SIZE) -> List[str]:
    """
    Разбивает текст на чанки по абзацам, не превышающие max_chars.
    """
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    chunks: List[str] = []
    current: List[str] = []
    current_len = 0

    for p in paragraphs:
        plen = len(p) + 1
        if current and current_len + plen > max_chars:
            chunks.append("\n".join(current))
            current = [p]
            current_len = len(p)
        else:
            current.append(p)
            current_len += plen

    if current:
        chunks.append("\n".join(current))

    return chunks


def _build_vector_index(session: Dict[str, Any]) -> None:
    """
    Строит векторный индекс для текста сессии (session["cache"]).
    Результат кладётся в session["chunks"] и session["embeddings"].
    """
    text = (session.get("cache") or "").strip()
    if not text:
        return

    chunks = _split_text_into_chunks(text, max_chars=CHUNK_SIZE)
    if not chunks:
        return

    try:
        vectors = _embed_texts(chunks)
    except Exception as e:
        print(f"❌ Ошибка при расчёте эмбеддингов: {e}")
        session["chunks"] = []
        session["embeddings"] = []
        return

    if len(vectors) != len(chunks):
        print("⚠️ Количество эмбеддингов не совпало с количеством чанков")
        session["chunks"] = []
        session["embeddings"] = []
        return

    session["chunks"] = chunks
    session["embeddings"] = vectors
    print(f"✅ Построен векторный индекс: {len(chunks)} чанков")


def _search_similar_chunks(question: str, session: Dict[str, Any], top_k: int = 8) -> List[Tuple[float, str]]:
    """
    Возвращает top_k (similarity, chunk_text) для текущей сессии.
    Если индекс отсутствует или пуст — [].
    """
    chunks = session.get("chunks") or []
    vectors = session.get("embeddings") or []
    if not chunks or not vectors:
        return []

    try:
        q_vec = _embed_texts([question])[0]
    except Exception as e:
        print(f"❌ Ошибка эмбеддинга вопроса: {e}")
        return []

    q = np.array(q_vec, dtype=float)
    q_norm = np.linalg.norm(q) + 1e-8

    sims: List[Tuple[float, int]] = []
    for idx, v in enumerate(vectors):
        v_arr = np.array(v, dtype=float)
        v_norm = np.linalg.norm(v_arr) + 1e-8
        sim = float(np.dot(q, v_arr) / (q_norm * v_norm))
        sims.append((sim, idx))

    sims.sort(key=lambda x: x[0], reverse=True)
    top = sims[:top_k]
    return [(sim, chunks[i]) for sim, i in top]


# ===================== API: СЕССИЯ =====================
@app.post("/api/session")
def create_or_get_session():
    data = request.get_json(silent=True) or {}
    session_id = data.get("session_id") or str(uuid.uuid4())
    _ensure_session(session_id)
    return jsonify({"session_id": session_id, "message": "Сессия готова"}), 200


# ===================== API: ЗАГРУЗКА ФАЙЛОВ =====================
@app.post("/api/upload_files")
def upload_files():
    session_id = request.form.get("session_id", "")
    session = _ensure_session(session_id)

    if "files" not in request.files:
        return jsonify({"error": "Файлы не получены"}), 400

    uploaded = request.files.getlist("files")
    if not uploaded:
        return jsonify({"error": "Список файлов пуст"}), 400

    base = Path(session["dir"])
    added_files = []
    for storage in uploaded:
        fname = storage.filename or "file"
        if not _is_allowed(fname):
            continue
        dest = base / fname
        storage.save(dest)
        if str(dest).lower().endswith(".zip"):
            try:
                with zipfile.ZipFile(dest, "r") as zf:
                    zf.extractall(base)
            except Exception as e:
                return jsonify({"error": f"Ошибка распаковки {fname}: {e}"}), 400
            finally:
                dest.unlink(missing_ok=True)
        else:
            size = dest.stat().st_size
            session["files"].append({"name": fname, "size": size})
            added_files.append((fname, str(dest)))

    # Собираем все файлы (включая из ZIP)
    for p in base.rglob("*"):
        if p.is_file() and _is_allowed(p.name) and not str(p).lower().endswith(".zip"):
            if not any(p.name == f["name"] for f in session["files"]):
                session["files"].append({"name": p.name, "size": p.stat().st_size})
                added_files.append((p.name, str(p)))

    # Ингестим в кэш
    for shown_name, file_path in added_files:
        _ingest_file_to_cache(session, file_path, shown_name)

    return jsonify({
        "session_id": session_id,
        "message": f"Загружено файлов: {len(added_files)}",
        "files": session["files"],
        "cache_size": len(session["cache"])
    }), 200


# ===================== API: ОЧИСТИТЬ СЕССИЮ =====================
@app.post("/api/clear")
def clear_session():
    data = request.get_json(silent=True) or {}
    session_id = data.get("session_id", "")
    if session_id in SESSIONS:
        base = Path(SESSIONS[session_id]["dir"])
        try:
            shutil.rmtree(base, ignore_errors=True)
        except Exception:
            pass
        del SESSIONS[session_id]
    return jsonify({"message": "Сессия очищена"}), 200


# ===================== API: ВОПРОС-ОТВЕТ (с RAG) =====================
@app.post("/api/submit_question")
def submit_question():
    data = request.get_json(silent=True) or {}
    session_id = data.get("session_id", "")
    question = (data.get("question") or "").strip()

    if not question:
        return jsonify({"error": "Введите вопрос"}), 400

    session = _ensure_session(session_id)
    context = session["cache"] or ""
    if not context.strip():
        return jsonify({"error": "Нет данных для поиска. Загрузите документы."}), 400

    headers = {"Authorization": f"Bearer {XAI_API_KEY}", "Content-Type": "application/json"}

    # 1) Убеждаемся, что есть векторный индекс. Если нет — строим.
    if not session.get("chunks") or not session.get("embeddings"):
        _build_vector_index(session)

    # 2) Векторный поиск релевантных чанков
    similar_chunks = _search_similar_chunks(question, session, top_k=8)

    if similar_chunks:
        # Можно ввести порог по похожести, чтобы отсечь откровенный шум
        filtered_texts = [text for sim, text in similar_chunks if sim > 0.15]
        if not filtered_texts:  # если всё ниже порога — берём как есть
            filtered_texts = [text for _, text in similar_chunks]
        retrieved_context = "\n\n---\n\n".join(filtered_texts)
    else:
        # fallback: если RAG не сработал, усечём общий контекст
        retrieved_context = context[:150000]

    prompt = (
        "Ниже приведены фрагменты загруженных пользователем документов "
        "(могут быть разные файлы и листы). "
        "Ответь на вопрос строго по этим фрагментам, без внешних знаний. "
        "По возможности указывай источники (FILE/лист/страница), если они упоминаются в тексте.\n\n"
        f"{retrieved_context}\n\n"
        f"Вопрос: {question}\n\n"
        "Если информации недостаточно — честно напиши «Информация отсутствует»."
    )

    try:
        resp = requests.post(API_URL, headers=headers, json={
            "model": MODEL,
            "messages": [
                {
                    "role": "system",
                    "content": "Ты эксперт по документам. Отвечай только по предоставленному контексту и по возможности указывай источники."
                },
                {"role": "user", "content": prompt}
            ],
            "max_output_tokens": OUTPUT_TOKENS
        })
        resp.raise_for_status()
        rj = resp.json()
        if "output" in rj:
            final_answer = rj["output"][0].get("content", "")
        else:
            final_answer = rj["choices"][0]["message"]["content"]
    except requests.exceptions.HTTPError as e:
        err = e.response.text if e.response is not None else str(e)
        final_answer = f"Ошибка при обращении к X.AI API: {err}"
    except Exception as e:
        final_answer = f"Ошибка обработки запроса: {e}"

    session["last_question"] = question
    session["last_answer"] = final_answer
    return jsonify({"answer": final_answer}), 200


# ===================== API: ЭКСПОРТ ОТВЕТА В XLSX =====================
@app.get("/api/export_xlsx")
def export_xlsx():
    session_id = request.args.get("session_id", "")
    if session_id not in SESSIONS:
        return jsonify({"error": "Сессия не найдена"}), 400

    question = SESSIONS[session_id].get("last_question", "").strip()
    answer = SESSIONS[session_id].get("last_answer", "").strip()
    if not question or not answer:
        return jsonify({"error": "Нет результата для выгрузки. Сначала выполните запрос."}), 400

    wb = Workbook()
    ws = wb.active
    ws.title = "Answer"

    ws["A1"] = "Вопрос"
    ws["A2"] = question
    ws["A4"] = "Ответ"
    ws["A5"] = answer

    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)

    fname = f"answer_{session_id}.xlsx"
    return send_file(
        bio,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=fname
    )


# ===================== СТАТИКА =====================
@app.route('/')
def serve_index():
    return send_from_directory('.', 'index.html')


@app.route('/favicon.ico')
def favicon():
    return '', 204


# ===================== ЗАПУСК =====================
if __name__ == '__main__':
    print("🚀 Запуск Flask-сервера на :5000")
    app.run(host='0.0.0.0', port=5000, debug=False)
