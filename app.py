# -*- coding: utf-8 -*-
"""
Локальный Flask-сервер для Q&A по загруженным пользователем документам:
- Загрузка файлов (PDF, DOCX, XLSX, TXT) в рамках пользовательской сессии
- Извлечение текста и формирование кэша
- Вопрос → поиск/анализ строго по кэшу
- Вывод ответа + экспорт результата в XLSX

Зависимости: flask, requests, python-docx, PyPDF2, openpyxl
"""
import os
import io
import json
import uuid
import shutil
import zipfile
from pathlib import Path
from typing import Dict, Any

from flask import Flask, request, jsonify, send_from_directory, send_file
import requests
from docx import Document
import PyPDF2
from openpyxl import load_workbook as xl_load
from openpyxl.workbook import Workbook

# ===================== НАСТРОЙКИ =====================
app = Flask(__name__, static_folder='.')
app.config["MAX_CONTENT_LENGTH"] = 512 * 1024 * 1024  # 512MB на всякий случай

# Ключи/модель как в примере: логика и движок не меняются
XAI_API_KEY = os.environ.get("XAI_API_KEY")
API_URL = "https://api.x.ai/v1/chat/completions"
MODEL = "grok-4-1-fast-reasoning-latest"
OUTPUT_TOKENS = 10000

# Чанкование кэша (как в примере: ~1900к символов на запрос)
CHUNK_SIZE = 1900000

# Сессии (in-memory). Значение: {"dir": str, "cache": str, "files": [ {name, size} ], "last_answer": str, "last_question": str}
SESSIONS: Dict[str, Dict[str, Any]] = {}

# ===================== УТИЛИТЫ =====================
def _ensure_session(session_id: str) -> Dict[str, Any]:
    if not session_id:
        session_id = str(uuid.uuid4())
    if session_id not in SESSIONS:
        base = Path("/tmp") / f"session_{session_id}"
        base.mkdir(parents=True, exist_ok=True)
        SESSIONS[session_id] = {
            "dir": str(base),
            "cache": "",
            "files": [],
            "last_answer": "",
            "last_question": ""
        }
    return SESSIONS[session_id]

def _is_allowed(fname: str) -> bool:
    low = fname.lower()
    return (
        low.endswith(".pdf")
        or low.endswith(".docx")
        or low.endswith(".xlsx")
        or low.endswith(".txt")      # ✅ добавлено
        or low.endswith(".zip")
    )

def _extract_text_from_pdf(path: str) -> str:
    try:
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            return "\n".join([page.extract_text() or "" for page in reader.pages])
    except Exception as e:
        return f"[PDF error: {e}]"

def _extract_text_from_docx(path: str) -> str:
    try:
        doc = Document(path)
        parts = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        # простая выборка из таблиц (если есть)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if row_text:
                    parts.append(" | ".join(row_text))
        return "\n".join(parts)
    except Exception as e:
        return f"[DOCX error: {e}]"

def _extract_text_from_xlsx(path: str) -> str:
    """
    Плоское извлечение текста из всех листов XLSX, по строкам.
    Формат: "SheetName: A1 | A2 | ...".
    """
    try:
        wb = xl_load(filename=path, read_only=True, data_only=True)
        out_lines = []
        for ws in wb.worksheets:
            out_lines.append(f"=== SHEET: {ws.title} ===")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip() != ""]
                if cells:
                    out_lines.append(" | ".join(cells))
        return "\n".join(out_lines)
    except Exception as e:
        return f"[XLSX error: {e}]"

def _extract_text_from_txt(path: str) -> str:
    """ ✅ Новая функция извлечения TXT """
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()
    except Exception as e:
        return f"[TXT error: {e}]"

def _ingest_file_to_cache(session: Dict[str, Any], file_path: str, shown_name: str) -> None:
    text = ""
    low = shown_name.lower()
    if low.endswith(".pdf"):
        text = _extract_text_from_pdf(file_path)
    elif low.endswith(".docx"):
        text = _extract_text_from_docx(file_path)
    elif low.endswith(".xlsx"):
        text = _extract_text_from_xlsx(file_path)
    elif low.endswith(".txt"):                 # ✅ добавлено
        text = _extract_text_from_txt(file_path)

    if text:
        # Вставляем "маркеры источника" для удобства ссылки в ответах
        header = f"\n\n### FILE: {shown_name}\n"
        session["cache"] += header + text

# ===================== API: СЕССИЯ =====================
@app.post("/api/session")
def create_or_get_session():
    data = request.get_json(silent=True) or {}
    session_id = data.get("session_id") or str(uuid.uuid4())
    _ensure_session(session_id)
    return jsonify({"session_id": session_id, "message": "Сессия готова"}), 200

# ===================== API: ЗАГРУЗКА ФАЙЛОВ =====================
@app.post("/api/upload_files")
def upload_files():
    session_id = request.form.get("session_id", "")
    session = _ensure_session(session_id)

    if "files" not in request.files:
        return jsonify({"error": "Файлы не получены"}), 400

    uploaded = request.files.getlist("files")
    if not uploaded:
        return jsonify({"error": "Список файлов пуст"}), 400

    base = Path(session["dir"])
    added_files = []
    for storage in uploaded:
        fname = storage.filename or "file"
        if not _is_allowed(fname):
            continue
        dest = base / fname
        storage.save(dest)
        if str(dest).lower().endswith(".zip"):
            try:
                with zipfile.ZipFile(dest, "r") as zf:
                    zf.extractall(base)
            except Exception as e:
                return jsonify({"error": f"Ошибка распаковки {fname}: {e}"}), 400
            finally:
                dest.unlink(missing_ok=True)
        else:
            size = dest.stat().st_size
            session["files"].append({"name": fname, "size": size})
            added_files.append((fname, str(dest)))

    for p in base.rglob("*"):
        if p.is_file() and _is_allowed(p.name) and not str(p).lower().endswith(".zip"):
            if not any(p.name == f["name"] for f in session["files"]):
                session["files"].append({"name": p.name, "size": p.stat().st_size})
                added_files.append((p.name, str(p)))

    for shown_name, file_path in added_files:
        _ingest_file_to_cache(session, file_path, shown_name)

    return jsonify({
        "session_id": session_id,
        "message": f"Загружено файлов: {len(added_files)}",
        "files": session["files"],
        "cache_size": len(session["cache"])
    }), 200

# ===================== API: ОЧИСТИТЬ СЕССИЮ =====================
@app.post("/api/clear")
def clear_session():
    data = request.get_json(silent=True) or {}
    session_id = data.get("session_id", "")
    if session_id in SESSIONS:
        base = Path(SESSIONS[session_id]["dir"])
        try:
            shutil.rmtree(base, ignore_errors=True)
        except Exception:
            pass
        del SESSIONS[session_id]
    return jsonify({"message": "Сессия очищена"}), 200

# ===================== API: ВОПРОС-ОТВЕТ =====================
@app.post("/api/submit_question")
def submit_question():
    data = request.get_json(silent=True) or {}
    session_id = data.get("session_id", "")
    question = (data.get("question") or "").strip()

    if not question:
        return jsonify({"error": "Введите вопрос"}), 400

    session = _ensure_session(session_id)
    context = session["cache"] or ""
    if not context.strip():
        return jsonify({"error": "Нет данных для поиска. Загрузите документы."}), 400

    num_chunks = (len(context) // CHUNK_SIZE) + 1
    headers = {"Authorization": f"Bearer {XAI_API_KEY}", "Content-Type": "application/json"}

    summary = ""
    for i in range(num_chunks):
        start = i * CHUNK_SIZE
        end = start + CHUNK_SIZE
        chunk = context[start:end]
        if not chunk:
            continue

        prompt = (
            f"Используя следующий контекст:\n\n{chunk}\n\n"
            f"Ответь на вопрос: {question}. Укажи источники (название файла/лист/страница, если это возможно). "
            f"Информация должна браться строго из контекста."
        )

        try:
            resp = requests.post(API_URL, headers=headers, json={
                "model": MODEL,
                "messages": [
                    {"role": "system", "content": "Ты эксперт по документам. Отвечай строго по контексту и ссылайся на источник (FILE/лист/страница)."},
                    {"role": "user", "content": prompt}
                ],
                "max_output_tokens": OUTPUT_TOKENS
            })
            resp.raise_for_status()
            rj = resp.json()
            if "output" in rj:
                part_answer = rj["output"][0].get("content", "")
            else:
                part_answer = rj["choices"][0]["message"]["content"]
            summary += f"\n\n=== Часть {i + 1}/{num_chunks} ===\n{part_answer}"
        except requests.exceptions.HTTPError as e:
            err = e.response.text if e.response is not None else str(e)
            summary += f"\n\n=== Часть {i + 1}/{num_chunks} ===\nОшибка X.AI API: {err}"
        except Exception as e:
            summary += f"\n\n=== Часть {i + 1}/{num_chunks} ===\nОшибка обработки: {e}"

    final_answer = ""
    try:
        final_prompt = (
            f"Собери единый связный ответ на вопрос ниже, основываясь ТОЛЬКО на фрагментах после линии. "
            f"Сохрани ссылки на источники (FILE/лист/страница), если они были. "
            f"Вопрос: {question}\n\n"
            f"=======================\n{summary}"
        )
        final_resp = requests.post(API_URL, headers=headers, json={
            "model": MODEL,
            "messages": [
                {"role": "system", "content": "Ты эксперт по документам. Объедини ответы, сохрани источники, не добавляй внешние знания."},
                {"role": "user", "content": final_prompt}
            ],
            "max_output_tokens": OUTPUT_TOKENS
        })
        final_resp.raise_for_status()
        final_answer = final_resp.json()["choices"][0]["message"]["content"]
    except requests.exceptions.HTTPError as e:
        err = e.response.text if e.response is not None else str(e)
        final_answer = f"Ошибка при финальной обработке: {err}\n\nПромежуточные ответы:\n{summary}"
    except Exception as e:
        final_answer = f"Ошибка при финальной обработке: {e}\n\nПромежуточные ответы:\n{summary}"

    session["last_question"] = question
    session["last_answer"] = final_answer
    return jsonify({"answer": final_answer}), 200

# ===================== API: ЭКСПОРТ ОТВЕТА В XLSX =====================
@app.get("/api/export_xlsx")
def export_xlsx():
    session_id = request.args.get("session_id", "")
    if session_id not in SESSIONS:
        return jsonify({"error": "Сессия не найдена"}), 400

    question = SESSIONS[session_id].get("last_question", "").strip()
    answer = SESSIONS[session_id].get("last_answer", "").strip()
    if not question or not answer:
        return jsonify({"error": "Нет результата для выгрузки. Сначала выполните запрос."}), 400

    wb = Workbook()
    ws = wb.active
    ws.title = "Answer"

    ws["A1"] = "Вопрос"
    ws["A2"] = question
    ws["A4"] = "Ответ"
    ws["A5"] = answer

    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)

    fname = f"answer_{session_id}.xlsx"
    return send_file(
        bio,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=fname
    )

# ===================== СТАТИКА =====================
@app.route('/')
def serve_index():
    return send_from_directory('.', 'index.html')

@app.route('/favicon.ico')
def favicon():
    return '', 204

# ===================== ЗАПУСК =====================
if __name__ == '__main__':
    print("🚀 Запуск Flask-сервера на :5000")
    app.run(host='0.0.0.0', port=5000, debug=False)
